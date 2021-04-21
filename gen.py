import docx
import yaml
import sys


def contar_temas(examen):
    c = 0
    for k in examen.keys():
        c += k.count('tema')
    return c


def buscar_encabezado(titulo, parrafo):
    i = 0
    for v in parrafo:
        if v.text.count(titulo):
            return i
        i += 1


def reemplazar_encabezado(old, new, parrafo):
    try:
        linea = buscar_encabezado(old, parrafo)
        palabra = buscar_encabezado(old, parrafo[linea].runs)
        texto = parrafo[linea].runs[palabra]
        texto.text = texto.text.replace(old, new)
    except TypeError:
        print('Hubo un problema al reemplazar "' + old + '" por "' + new +
              '". Puede ser un problema con el formato del archivo "' + template + '".')


def set_lang(lang, doc):
    # For new document (document-wide):
    # Set language value in the documents' default Run's Properties element.
    styles_element = doc.styles.element
    rpr_default = styles_element.xpath('./w:docDefaults/w:rPrDefault/w:rPr')[0]
    lang_default = rpr_default.xpath('w:lang')[0]
    lang_default.set(docx.oxml.shared.qn('w:val'), lang)


archivo = sys.argv[1]
out_p = archivo
out_p = out_p[:out_p.rfind('\\')]
# archivo = 'Otros/test2.yaml'

f = open(archivo, 'r', encoding='utf-8')
examen = yaml.safe_load(f)
f.close()

# Selección de template docx
template = 'Templates/base_' + \
    examen['encabezado']['colegio'].lower() + '.docx'
document = docx.Document(template)

# Idioma del documento
set_lang('es-ES', document)

# Escribe el encabezado
if examen.get('encabezado', 0):
    print('Escribiendo encabezado...')
    parrafos = document.paragraphs
    asg, prof, fecha, curso, ptos, tipoe, colegio = examen['encabezado'].values(
    )

    reemplazar_encabezado('asignatura', asg, parrafos)
    reemplazar_encabezado('prof', prof, parrafos)
    reemplazar_encabezado('fecha', fecha, parrafos)
    reemplazar_encabezado('curso', curso, parrafos)
    reemplazar_encabezado('puntos', ptos, parrafos)
    reemplazar_encabezado('tipo', tipoe, parrafos)

else:
    print('No se ha definido un encabezado. No se generará ningún documento.')
    sys.exit()

# Escribir el contenido
if contar_temas(examen):
    print("Escribiendo temas...")

    for k, v in examen.items():

        if k.startswith('tema '):
            # Escribe el título de cada tema
            titulo = k.capitalize() + ': ' + v['descripcion']
            document.add_heading(titulo, level=2)
            # Escribe los subtemas
            for subtema in v['subtemas']:
                if subtema['tipo'] == 'sel-m':
                    document.add_paragraph(
                        subtema['enunciado'], style='Normal')
                    for opcion in subtema['opciones']:
                        document.add_paragraph(
                            str(opcion), style='List Bullet')

                elif subtema['tipo'] == 'v-f':
                    for item in subtema['items']:
                        document.add_paragraph(
                            item + ' [V] [F]\n\n', style='List Bullet')

                elif subtema['tipo'] == 'cita':
                    document.add_paragraph(
                        subtema['enunciado'], style='Normal')
                    aux = document.add_paragraph()
                    tabs = 15 // subtema['cantidad']
                    for k in range(subtema['cantidad']):
                        if subtema['cantidad'] <= 3:
                            aux.add_run(str(k+1) + ' -' + '\t'*tabs)
                        else:
                            if k != subtema['cantidad'] - 1:
                                aux.add_run(str(k+1) + ' -' + '\n')
                            else:
                                aux.add_run(str(k+1) + ' -')

else:
    print('No se definió ningún tema. No se generará ningún documento.')
    sys.exit()

nombre = asg + ', ' + curso + ', ' + tipoe + ' - (' + colegio.upper() + ')'
nombre = out_p + '\\' + nombre + '.docx'
document.save(nombre)
print('Documento generado en "' + nombre + '".\nPresiona ENTER para salir.')
input()